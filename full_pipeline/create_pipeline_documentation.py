#!/usr/bin/env python3
"""
Create comprehensive documentation for the Hierarchical Monthly Route Pipeline
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_documentation():
    """Create comprehensive Word document explaining the pipeline"""

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading('Hierarchical Monthly Route Plan Pipeline', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Complete Processing Flow and Scenario Handling')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(128, 128, 128)

    # Date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================
    # SECTION 1: OVERVIEW
    # ==========================
    doc.add_heading('1. Pipeline Overview', 1)

    doc.add_paragraph(
        'The Hierarchical Monthly Route Plan Pipeline is an optimization system that processes '
        'route plans for sales agents, optimizing customer visit sequences using TSP (Traveling '
        'Salesman Problem) algorithms while maintaining a hierarchical processing structure.'
    )

    doc.add_heading('1.1 Key Objectives', 2)
    objectives = [
        'Optimize customer visit sequences to minimize travel distance',
        'Process route plans hierarchically by Distributor → Agent → Date',
        'Handle multiple scenarios based on customer counts',
        'Integrate prospective customers to reach optimal route density',
        'Assign sequential stop numbers per date for each agent',
        'Maintain data integrity across customer and prospect databases'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('1.2 Processing Hierarchy', 2)
    doc.add_paragraph('The pipeline follows a strict three-level hierarchy:')

    hierarchy_table = doc.add_table(rows=4, cols=3)
    hierarchy_table.style = 'Light Grid Accent 1'

    # Header
    header_cells = hierarchy_table.rows[0].cells
    header_cells[0].text = 'Level'
    header_cells[1].text = 'Description'
    header_cells[2].text = 'Processing Order'

    # Data
    data_rows = [
        ('1. DistributorID', 'Top-level grouping by distributor', 'Sequential'),
        ('2. AgentID', 'Sales agents within distributor', 'Sequential'),
        ('3. RouteDate', 'Individual dates per agent', 'Chronological (ASC)')
    ]

    for i, (level, desc, order) in enumerate(data_rows, start=1):
        cells = hierarchy_table.rows[i].cells
        cells[0].text = level
        cells[1].text = desc
        cells[2].text = order

    doc.add_page_break()

    # ==========================
    # SECTION 2: STEP-BY-STEP PROCESS
    # ==========================
    doc.add_heading('2. Step-by-Step Processing Flow', 1)

    # Step 1
    doc.add_heading('Step 1: Initialize Pipeline', 2)
    doc.add_paragraph(
        'The pipeline is initialized with configuration parameters:'
    )
    params = [
        'batch_size: Number of records to process in batches (default: 50)',
        'max_workers: Maximum concurrent workers for parallel processing (default: 4)',
        'start_lat/start_lon: Optional starting coordinates for TSP optimization',
        'Logging configuration: Creates timestamped log files in logs/ directory'
    ]
    for param in params:
        doc.add_paragraph(param, style='List Bullet 2')

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:32-51')

    # Step 2
    doc.add_heading('Step 2: Build Hierarchical Structure', 2)
    doc.add_paragraph(
        'The system queries MonthlyRoutePlan_temp to build the complete hierarchy:'
    )

    doc.add_paragraph('2.1 Extract All Distributors', style='List Number')
    doc.add_paragraph(
        'Query: SELECT DISTINCT DistributorID FROM MonthlyRoutePlan_temp',
        style='List Bullet 2'
    )

    doc.add_paragraph('2.2 For Each Distributor, Get All Agents', style='List Number')
    doc.add_paragraph(
        'Query: Filter by DistributorID and get all AgentIDs',
        style='List Bullet 2'
    )

    doc.add_paragraph('2.3 For Each Agent, Get All Dates (Chronological)', style='List Number')
    doc.add_paragraph(
        'Query: Get RouteDate, customer count, and total records',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'ORDER BY RouteDate ASC to ensure chronological processing',
        style='List Bullet 2'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:70-144')

    # Step 3
    doc.add_heading('Step 3: Process Each Agent with Sequential StopNo', 2)
    doc.add_paragraph(
        'This is the core processing step where routes are optimized for each agent across all their dates.'
    )

    doc.add_paragraph('3.1 Sort Dates Chronologically', style='List Number')
    doc.add_paragraph(
        'Ensures dates are processed in order (earliest to latest)',
        style='List Bullet 2'
    )

    doc.add_paragraph('3.2 Initialize Sequential Numbering', style='List Number')
    doc.add_paragraph(
        'StopNo starts fresh per date (1, 2, 3... N)',
        style='List Bullet 2'
    )

    doc.add_paragraph('3.3 Process Each Date', style='List Number')
    doc.add_paragraph(
        'For each date, perform scenario checking and data enrichment',
        style='List Bullet 2'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:146-408')

    # Step 4
    doc.add_heading('Step 4: Scenario Condition Checking', 2)
    doc.add_paragraph(
        'Before processing each date, the system checks which scenario applies based on customer count.'
    )

    scenario_table = doc.add_table(rows=5, cols=4)
    scenario_table.style = 'Light Grid Accent 1'

    # Header
    header_cells = scenario_table.rows[0].cells
    header_cells[0].text = 'Scenario'
    header_cells[1].text = 'Customer Count'
    header_cells[2].text = 'Action'
    header_cells[3].text = 'Processing'

    # Data
    scenario_data = [
        ('High Volume', '25+', 'Process with full optimization', 'Enabled'),
        ('Medium Volume', '10-24', 'Process with standard optimization', 'Enabled'),
        ('Low Volume', '5-9', 'Process with light optimization', 'Enabled'),
        ('Very Small', '1-4', 'Process minimal routes', 'Enabled')
    ]

    for i, (scenario, count, action, processing) in enumerate(scenario_data, start=1):
        cells = scenario_table.rows[i].cells
        cells[0].text = scenario
        cells[1].text = count
        cells[2].text = action
        cells[3].text = processing

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Currently, all scenarios are processed. The scenario classification helps with '
        'logging and potential future optimizations based on route size.'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:409-455')

    doc.add_page_break()

    # Step 5
    doc.add_heading('Step 5: Data Enrichment', 2)
    doc.add_paragraph(
        'For each date that passes scenario checks, the system enriches the data with coordinates '
        'and additional customer information.'
    )

    doc.add_paragraph('5.1 Fetch Monthly Plan Data', style='List Number')
    doc.add_paragraph(
        'Get all customer records from MonthlyRoutePlan_temp for the specific '
        'Distributor/Agent/Date combination',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'NOTE: Existing StopNo values are IGNORED at this stage',
        style='List Bullet 2'
    )

    doc.add_paragraph('5.2 Retrieve Customer Coordinates', style='List Number')
    doc.add_paragraph(
        'Join with customer table to get latitude, longitude, and barangay_code (address3)',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Filter out invalid coordinates (NULL, 0.0, or near-zero values)',
        style='List Bullet 2'
    )

    doc.add_paragraph('5.3 Separate Customers', style='List Number')
    doc.add_paragraph(
        'With coordinates: Ready for TSP optimization',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Without coordinates: Will receive StopNo = 100',
        style='List Bullet 2'
    )

    doc.add_paragraph('5.4 Detect Customer Type (custype)', style='List Number')
    doc.add_paragraph(
        'Check if CustNo exists in customer table → custype = "customer"',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Check if CustNo exists in prospective table → custype = "prospect"',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Not found in either → custype = "unknown"',
        style='List Bullet 2'
    )

    doc.add_paragraph('5.5 Smart Prospect Addition', style='List Number')
    doc.add_paragraph(
        'Target: Reach 60 total customers per date (existing + prospects)',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Only add prospects from the same barangay_code as existing customers',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Exclude prospects that already have visits in custvisit table',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Exclude prospects already in MonthlyRoutePlan_temp for this date',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Randomized selection (ORDER BY NEWID())',
        style='List Bullet 2'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:531-761')

    # Step 6
    doc.add_heading('Step 6: TSP Optimization', 2)
    doc.add_paragraph(
        'Apply Traveling Salesman Problem optimization using the Nearest Neighbor heuristic '
        'to minimize travel distance.'
    )

    doc.add_paragraph('6.1 Algorithm: Nearest Neighbor with Haversine Distance', style='List Number')
    doc.add_paragraph(
        'Start from specified location (if provided) or first customer',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Iteratively select the nearest unvisited customer',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Use Haversine formula to calculate great-circle distance',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Continue until all customers are visited',
        style='List Bullet 2'
    )

    doc.add_paragraph('6.2 Haversine Distance Formula', style='List Number')
    doc.add_paragraph(
        'Calculates shortest distance between two points on Earth\'s surface',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Accounts for Earth\'s curvature using latitude/longitude',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Returns distance in kilometers',
        style='List Bullet 2'
    )

    doc.add_paragraph('6.3 Output', style='List Number')
    doc.add_paragraph(
        'Optimized DataFrame with temporary stopno (1, 2, 3... N)',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Maintains all original customer data',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'RouteDate preserved for tracking',
        style='List Bullet 2'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:457-529')

    doc.add_page_break()

    # Step 7
    doc.add_heading('Step 7: Per-Date Sequential StopNo Assignment', 2)
    doc.add_paragraph(
        'After TSP optimization for all dates, assign fresh sequential stop numbers for each date.'
    )

    doc.add_paragraph('7.1 Reset Numbering Per Date', style='List Number')
    doc.add_paragraph(
        'For each RouteDate, start StopNo from 1',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'This ensures consistent numbering across all agents and dates',
        style='List Bullet 2'
    )

    doc.add_paragraph('7.2 Assign StopNo to Optimized Customers', style='List Number')
    doc.add_paragraph(
        'Customers with coordinates: StopNo = 1, 2, 3... N (in optimized order)',
        style='List Bullet 2'
    )

    doc.add_paragraph('7.3 Assign StopNo to Customers Without Coordinates', style='List Number')
    doc.add_paragraph(
        'Customers without coordinates: StopNo = 100',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Ensures they are processed but not in the optimized sequence',
        style='List Bullet 2'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:199-246')

    # Step 8
    doc.add_heading('Step 8: Database Updates and Inserts', 2)
    doc.add_paragraph(
        'Update existing customer records and insert new prospect records into MonthlyRoutePlan_temp.'
    )

    doc.add_paragraph('8.1 Separate Updates from Inserts', style='List Number')
    doc.add_paragraph(
        'Existing customers (custype = "customer"): UPDATE StopNo',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'New prospects (custype = "prospect"): INSERT with all required fields',
        style='List Bullet 2'
    )

    doc.add_paragraph('8.2 Batch Update Existing Customers', style='List Number')
    doc.add_paragraph(
        'Use executemany() for efficient batch updates',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Update query: SET StopNo WHERE DistributorID + AgentID + RouteDate + CustNo match',
        style='List Bullet 2'
    )

    doc.add_paragraph('8.3 Batch Insert Prospects', style='List Number')
    doc.add_paragraph(
        'Insert fields: DistributorID, AgentID, RouteDate, CustNo, StopNo, Name, WD, etc.',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Copy default values (WD, SalesManTerritory, etc.) from existing customers',
        style='List Bullet 2'
    )

    doc.add_paragraph('8.4 Transaction Management', style='List Number')
    doc.add_paragraph(
        'All operations within a transaction',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Commit on success, rollback on error',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Maintains data integrity',
        style='List Bullet 2'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:247-376')

    # Step 9
    doc.add_heading('Step 9: Update Customer Type Classification', 2)
    doc.add_paragraph(
        'After all route processing is complete, update the custype field using database JOINs.'
    )

    doc.add_paragraph('9.1 Update Customers', style='List Number')
    doc.add_paragraph(
        'JOIN MonthlyRoutePlan_temp with customer table',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'SET custype = "customer" where match found',
        style='List Bullet 2'
    )

    doc.add_paragraph('9.2 Update Prospects', style='List Number')
    doc.add_paragraph(
        'JOIN MonthlyRoutePlan_temp with prospective table',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'SET custype = "prospect" where match found',
        style='List Bullet 2'
    )

    doc.add_paragraph('9.3 Verify Completion', style='List Number')
    doc.add_paragraph(
        'Check for any records with NULL or "unknown" custype',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Log warnings for unknown records',
        style='List Bullet 2'
    )

    doc.add_paragraph('File: run_monthly_route_pipeline_hierarchical.py:884-934')

    doc.add_page_break()

    # ==========================
    # SECTION 3: SCENARIO HANDLING
    # ==========================
    doc.add_heading('3. Detailed Scenario Handling', 1)

    doc.add_heading('3.1 Scenario Classification Logic', 2)

    doc.add_paragraph(
        'The pipeline classifies each route into one of four scenarios based on customer count. '
        'This classification helps with logging, performance tracking, and potential future '
        'optimizations.'
    )

    # Detailed scenario table
    detailed_scenario_table = doc.add_table(rows=5, cols=2)
    detailed_scenario_table.style = 'Medium Grid 3 Accent 1'

    # Header
    header_cells = detailed_scenario_table.rows[0].cells
    header_cells[0].text = 'Scenario'
    header_cells[1].text = 'Description & Processing'

    # Data with details
    scenario_details = [
        ('High Volume\n(25+ customers)',
         'Routes with 25 or more customers. These routes benefit most from TSP optimization '
         'as distance savings compound with more stops. Full optimization applied with '
         'prospect addition to reach 60 customers.'),

        ('Medium Volume\n(10-24 customers)',
         'Standard routes with moderate customer counts. Good candidates for optimization '
         'and prospect addition. Most common scenario in typical operations.'),

        ('Low Volume\n(5-9 customers)',
         'Smaller routes that still benefit from optimization. TSP provides modest distance '
         'improvements. Prospect addition brings significant value by increasing route density.'),

        ('Very Small\n(1-4 customers)',
         'Minimal routes with very few customers. Limited benefit from TSP, but prospect '
         'addition can substantially improve route efficiency.')
    ]

    for i, (scenario, details) in enumerate(scenario_details, start=1):
        cells = detailed_scenario_table.rows[i].cells
        cells[0].text = scenario
        cells[1].text = details

    doc.add_heading('3.2 Scenario Processing Decision', 2)
    doc.add_paragraph(
        'Current Implementation: All scenarios are processed (should_process = True)'
    )
    doc.add_paragraph(
        'The scenario information is logged and tracked but does not affect whether a route '
        'is processed or skipped.'
    )

    doc.add_heading('3.3 Future Scenario-Based Optimizations', 2)
    doc.add_paragraph(
        'Potential future enhancements based on scenario:'
    )
    future_opts = [
        'High Volume: Use more sophisticated TSP algorithms (2-opt, genetic algorithms)',
        'Medium Volume: Standard nearest neighbor with local optimization',
        'Low Volume: Simplified optimization with reduced computation',
        'Very Small: Skip TSP, focus on prospect addition only'
    ]
    for opt in future_opts:
        doc.add_paragraph(opt, style='List Bullet')

    doc.add_page_break()

    # ==========================
    # SECTION 4: TECHNICAL DETAILS
    # ==========================
    doc.add_heading('4. Technical Implementation Details', 1)

    doc.add_heading('4.1 Database Schema', 2)
    doc.add_paragraph('Primary Table: MonthlyRoutePlan_temp')

    schema_table = doc.add_table(rows=12, cols=3)
    schema_table.style = 'Light Grid Accent 1'

    # Header
    header_cells = schema_table.rows[0].cells
    header_cells[0].text = 'Column'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Description'

    # Schema data
    schema_data = [
        ('DistributorID', 'VARCHAR', 'Distributor identifier'),
        ('AgentID', 'VARCHAR', 'Sales agent identifier'),
        ('RouteDate', 'DATE', 'Scheduled visit date'),
        ('CustNo', 'VARCHAR', 'Customer/prospect number'),
        ('StopNo', 'INT', 'Sequential stop number (optimized)'),
        ('Name', 'VARCHAR', 'Customer/prospect name'),
        ('WD', 'INT', 'Week day indicator'),
        ('SalesManTerritory', 'VARCHAR', 'Territory code'),
        ('RouteName', 'VARCHAR', 'Route name'),
        ('RouteCode', 'VARCHAR', 'Route code'),
        ('SalesOfficeID', 'VARCHAR', 'Sales office identifier')
    ]

    for i, (col, type_, desc) in enumerate(schema_data, start=1):
        cells = schema_table.rows[i].cells
        cells[0].text = col
        cells[1].text = type_
        cells[2].text = desc

    doc.add_heading('4.2 Key Algorithms', 2)

    doc.add_paragraph('Nearest Neighbor TSP Heuristic:', style='Heading 3')
    doc.add_paragraph(
        'Time Complexity: O(n²) where n is the number of customers'
    )
    doc.add_paragraph(
        'Space Complexity: O(n) for storing the route'
    )
    doc.add_paragraph(
        'Advantages: Fast, simple, produces reasonable solutions for most cases'
    )
    doc.add_paragraph(
        'Limitations: Not guaranteed to find optimal solution, can be 25-50% longer than optimal'
    )

    doc.add_paragraph('Haversine Distance Calculation:', style='Heading 3')
    doc.add_paragraph(
        'Formula: d = 2r × arcsin(√(sin²(Δφ/2) + cos(φ1) × cos(φ2) × sin²(Δλ/2)))'
    )
    doc.add_paragraph(
        'Where: φ = latitude, λ = longitude, r = Earth radius (6371 km)'
    )
    doc.add_paragraph(
        'Accuracy: Assumes spherical Earth (accurate within 0.5% for most purposes)'
    )

    doc.add_heading('4.3 Performance Characteristics', 2)

    perf_table = doc.add_table(rows=5, cols=2)
    perf_table.style = 'Light Grid Accent 1'

    header_cells = perf_table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value / Description'

    perf_data = [
        ('Processing Mode', 'Sequential (can enable parallel with --parallel flag)'),
        ('Batch Size', '50 records (configurable via --batch-size)'),
        ('Max Workers', '4 threads (configurable via --max-workers)'),
        ('Typical Rate', '5-20 combinations/second (depends on data complexity)')
    ]

    for i, (metric, value) in enumerate(perf_data, start=1):
        cells = perf_table.rows[i].cells
        cells[0].text = metric
        cells[1].text = value

    doc.add_heading('4.4 Error Handling', 2)
    error_handling = [
        'Transaction rollback on database errors',
        'Individual combination errors logged but don\'t stop pipeline',
        'Comprehensive logging to timestamped log files',
        'Graceful handling of missing coordinates',
        'Validation of stop number assignments',
        'Connection cleanup in finally blocks'
    ]
    for item in error_handling:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ==========================
    # SECTION 5: LOGGING AND MONITORING
    # ==========================
    doc.add_heading('5. Logging and Monitoring', 1)

    doc.add_heading('5.1 Log File Structure', 2)
    doc.add_paragraph(
        'Log files are created in: full_pipeline/logs/'
    )
    doc.add_paragraph(
        'Filename format: hierarchical_monthly_route_pipeline_YYYYMMDD_HHMMSS.log'
    )
    doc.add_paragraph(
        'Output: Both file and console (StreamHandler)'
    )

    doc.add_heading('5.2 Log Levels and Content', 2)

    log_table = doc.add_table(rows=5, cols=2)
    log_table.style = 'Light Grid Accent 1'

    header_cells = log_table.rows[0].cells
    header_cells[0].text = 'Level'
    header_cells[1].text = 'Content'

    log_data = [
        ('INFO', 'Processing progress, hierarchy structure, scenario classification, '
                 'customer counts, optimization results, update counts'),
        ('WARNING', 'Missing data, invalid coordinates, unknown customer types, '
                    'no prospects found, missing barangay codes'),
        ('ERROR', 'Database errors, query failures, optimization errors, '
                  'transaction rollbacks'),
        ('DEBUG', 'Not currently used (can be enabled for development)')
    ]

    for i, (level, content) in enumerate(log_data, start=1):
        cells = log_table.rows[i].cells
        cells[0].text = level
        cells[1].text = content

    doc.add_heading('5.3 Progress Tracking', 2)
    doc.add_paragraph('The pipeline provides real-time progress updates:')
    progress_items = [
        'Total combinations to process',
        'Current distributor/agent/date being processed',
        'Scenario classification for each route',
        'Customer counts (with/without coordinates)',
        'Prospect addition results',
        'TSP optimization progress',
        'Database update/insert counts',
        'Overall progress percentage',
        'Processing rate (combinations/second)'
    ]
    for item in progress_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.4 Final Summary Report', 2)
    doc.add_paragraph('At completion, the pipeline generates a summary:')
    summary_items = [
        'Total combinations processed',
        'Successful processing count',
        'Error count',
        'Skipped count',
        'Total duration',
        'Processing rate',
        'Overall status'
    ]
    for item in summary_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ==========================
    # SECTION 6: USAGE EXAMPLES
    # ==========================
    doc.add_heading('6. Usage Examples', 1)

    doc.add_heading('6.1 Basic Usage', 2)
    doc.add_paragraph('Run with default settings:')
    doc.add_paragraph(
        'python run_monthly_route_pipeline_hierarchical.py',
        style='Intense Quote'
    )

    doc.add_heading('6.2 Custom Batch Size', 2)
    doc.add_paragraph('Process with larger batches:')
    doc.add_paragraph(
        'python run_monthly_route_pipeline_hierarchical.py --batch-size 100',
        style='Intense Quote'
    )

    doc.add_heading('6.3 Parallel Processing', 2)
    doc.add_paragraph('Enable parallel processing with multiple workers:')
    doc.add_paragraph(
        'python run_monthly_route_pipeline_hierarchical.py --parallel --max-workers 8',
        style='Intense Quote'
    )

    doc.add_heading('6.4 Test Mode', 2)
    doc.add_paragraph('Run in test mode (limited processing):')
    doc.add_paragraph(
        'python run_monthly_route_pipeline_hierarchical.py --test-mode',
        style='Intense Quote'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 7: DATA FLOW DIAGRAM
    # ==========================
    doc.add_heading('7. Complete Data Flow', 1)

    doc.add_paragraph('The following describes the complete data flow through the system:')

    # Create a numbered list for flow
    flow_steps = [
        ('Input', 'MonthlyRoutePlan_temp table with initial route assignments'),
        ('Hierarchy Build', 'Extract Distributor → Agent → Date structure'),
        ('Iteration Start', 'Begin processing each distributor'),
        ('Agent Processing', 'For each agent, collect all dates chronologically'),
        ('Date Processing', 'For each date, check scenario conditions'),
        ('Data Enrichment', 'Join with customer table for coordinates'),
        ('Prospect Addition', 'Query prospective table for additional customers'),
        ('TSP Optimization', 'Apply nearest neighbor algorithm to locations'),
        ('StopNo Assignment', 'Assign sequential numbers per date'),
        ('Database Update', 'UPDATE existing customers with new StopNo'),
        ('Database Insert', 'INSERT new prospects into route plan'),
        ('Next Iteration', 'Move to next date/agent/distributor'),
        ('Final Update', 'Update custype using JOIN operations'),
        ('Output', 'Optimized MonthlyRoutePlan_temp with sequential stop numbers')
    ]

    for i, (stage, description) in enumerate(flow_steps, start=1):
        para = doc.add_paragraph(f'{stage}: ', style='List Number')
        para.add_run(description)

    doc.add_page_break()

    # ==========================
    # SECTION 8: BEST PRACTICES
    # ==========================
    doc.add_heading('8. Best Practices and Recommendations', 1)

    doc.add_heading('8.1 Data Preparation', 2)
    prep_items = [
        'Ensure customer table has accurate latitude/longitude for all active customers',
        'Populate address3 (barangay_code) for geographic grouping',
        'Clean up invalid coordinates (0.0, NULL, near-zero values)',
        'Verify prospective table has valid coordinates and barangay codes',
        'Check MonthlyRoutePlan_temp has all required fields populated'
    ]
    for item in prep_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.2 Performance Optimization', 2)
    perf_items = [
        'Run during off-peak hours to minimize database load',
        'Use --parallel flag for large datasets (1000+ combinations)',
        'Adjust --batch-size based on available memory',
        'Monitor log files for bottlenecks',
        'Consider database indexing on DistributorID, AgentID, RouteDate, CustNo'
    ]
    for item in perf_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.3 Data Quality', 2)
    quality_items = [
        'Regularly audit custype assignments',
        'Verify prospect additions are appropriate for each route',
        'Check for customers consistently receiving StopNo = 100',
        'Monitor scenario distribution to ensure balanced processing',
        'Review TSP optimization results for obviously suboptimal routes'
    ]
    for item in quality_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.4 Troubleshooting', 2)

    troubleshoot_table = doc.add_table(rows=6, cols=2)
    troubleshoot_table.style = 'Medium Grid 3 Accent 1'

    header_cells = troubleshoot_table.rows[0].cells
    header_cells[0].text = 'Issue'
    header_cells[1].text = 'Solution'

    troubleshoot_data = [
        ('No customers with coordinates',
         'Check customer table latitude/longitude population. Verify address3 exists for '
         'prospect matching.'),

        ('High error count',
         'Review log file for specific errors. Check database connection stability. '
         'Verify data integrity.'),

        ('Slow processing',
         'Reduce batch size, increase max workers, enable parallel mode. Check database '
         'query performance.'),

        ('Incorrect StopNo assignments',
         'Verify TSP optimization is running. Check for coordinate data quality issues. '
         'Review log for optimization errors.'),

        ('Missing prospects',
         'Ensure prospective table has data in same barangays as customers. Check NOT '
         'EXISTS conditions are not too restrictive.')
    ]

    for i, (issue, solution) in enumerate(troubleshoot_data, start=1):
        cells = troubleshoot_table.rows[i].cells
        cells[0].text = issue
        cells[1].text = solution

    doc.add_page_break()

    # ==========================
    # SECTION 9: APPENDIX
    # ==========================
    doc.add_heading('9. Appendix', 1)

    doc.add_heading('9.1 Key Database Queries', 2)

    doc.add_paragraph('Get Hierarchy Structure:', style='Heading 3')
    query1 = """SELECT DISTINCT DistributorID
FROM MonthlyRoutePlan_temp
WHERE DistributorID IS NOT NULL
ORDER BY DistributorID"""
    doc.add_paragraph(query1, style='Intense Quote')

    doc.add_paragraph('Get Customer Coordinates:', style='Heading 3')
    query2 = """SELECT CustNo, latitude, longitude, address3 as barangay_code
FROM customer
WHERE latitude IS NOT NULL
  AND longitude IS NOT NULL
  AND latitude != 0.0
  AND longitude != 0.0"""
    doc.add_paragraph(query2, style='Intense Quote')

    doc.add_paragraph('Get Prospects by Barangay:', style='Heading 3')
    query3 = """SELECT TOP N CustNo, Latitude, Longitude, barangay_code, OutletName
FROM prospective
WHERE barangay_code IN (...)
  AND Latitude IS NOT NULL
  AND Longitude IS NOT NULL
  AND NOT EXISTS (SELECT 1 FROM MonthlyRoutePlan_temp WHERE ...)
  AND NOT EXISTS (SELECT 1 FROM custvisit WHERE ...)
ORDER BY NEWID()"""
    doc.add_paragraph(query3, style='Intense Quote')

    doc.add_heading('9.2 Command Line Arguments', 2)

    args_table = doc.add_table(rows=5, cols=3)
    args_table.style = 'Light Grid Accent 1'

    header_cells = args_table.rows[0].cells
    header_cells[0].text = 'Argument'
    header_cells[1].text = 'Default'
    header_cells[2].text = 'Description'

    args_data = [
        ('--batch-size', '50', 'Number of records per batch'),
        ('--max-workers', '4', 'Maximum concurrent threads'),
        ('--parallel', 'False', 'Enable parallel processing'),
        ('--test-mode', 'False', 'Run in test mode (limited data)')
    ]

    for i, (arg, default, desc) in enumerate(args_data, start=1):
        cells = args_table.rows[i].cells
        cells[0].text = arg
        cells[1].text = default
        cells[2].text = desc

    doc.add_heading('9.3 File References', 2)
    file_refs = [
        'Main script: run_monthly_route_pipeline_hierarchical.py',
        'Database module: core/database.py',
        'Log directory: full_pipeline/logs/',
        'Configuration: Command line arguments'
    ]
    for ref in file_refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('9.4 Related Tables', 2)

    tables_table = doc.add_table(rows=4, cols=2)
    tables_table.style = 'Light Grid Accent 1'

    header_cells = tables_table.rows[0].cells
    header_cells[0].text = 'Table'
    header_cells[1].text = 'Purpose'

    tables_data = [
        ('MonthlyRoutePlan_temp', 'Main route plan table (input/output)'),
        ('customer', 'Customer master data with coordinates'),
        ('prospective', 'Prospective customer data for route expansion')
    ]

    for i, (table, purpose) in enumerate(tables_data, start=1):
        cells = tables_table.rows[i].cells
        cells[0].text = table
        cells[1].text = purpose

    # ==========================
    # FOOTER
    # ==========================
    doc.add_page_break()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('End of Documentation')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run2 = footer_para2.add_run(
        f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
    )
    footer_run2.font.size = Pt(9)
    footer_run2.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = r'C:\Simplr projects\Route-optimization\full_pipeline\Hierarchical_Monthly_Route_Pipeline_Documentation.docx'
    doc.save(output_path)

    print(f"Documentation created successfully!")
    print(f"Saved to: {output_path}")
    print(f"\nDocument sections:")
    print("  1. Pipeline Overview")
    print("  2. Step-by-Step Processing Flow (Steps 1-9)")
    print("  3. Detailed Scenario Handling")
    print("  4. Technical Implementation Details")
    print("  5. Logging and Monitoring")
    print("  6. Usage Examples")
    print("  7. Complete Data Flow")
    print("  8. Best Practices and Recommendations")
    print("  9. Appendix (Queries, Arguments, References)")

if __name__ == "__main__":
    create_documentation()
