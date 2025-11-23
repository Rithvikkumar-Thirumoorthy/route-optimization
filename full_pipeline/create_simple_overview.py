#!/usr/bin/env python3
"""
Create a simple, easy-to-understand overview document for non-technical users
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_simple_overview():
    """Create simple overview document for non-technical audience"""

    doc = Document()

    # Title
    title = doc.add_heading('Route Optimization System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('A Simple Guide for Everyone')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.color.rgb = RGBColor(0, 128, 0)

    intro = doc.add_paragraph(
        'This guide explains how our system helps sales agents visit customers '
        'in the most efficient order, saving time and fuel.'
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f'{datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================
    # SECTION 1: THE PROBLEM
    # ==========================
    doc.add_heading('1. The Problem We\'re Solving', 1)

    doc.add_heading('Imagine This Situation:', 2)

    problem_para = doc.add_paragraph()
    problem_para.add_run('A sales agent has 20 customers to visit tomorrow.\n').bold = True
    problem_para.add_run(
        'These customers are scattered around town. Without planning, the agent might:\n'
        '• Drive back and forth across town multiple times\n'
        '• Waste hours in traffic\n'
        '• Use more fuel than necessary\n'
        '• Visit fewer customers in a day\n'
        '• Feel exhausted from unnecessary travel'
    )

    doc.add_paragraph()

    doc.add_heading('What We Need:', 2)
    solution_items = [
        'Find the shortest route that visits all 20 customers',
        'Give each customer a visit number (1st, 2nd, 3rd... 20th)',
        'Help the agent follow an efficient path',
        'Add new potential customers along the way (if there\'s time)'
    ]
    for item in solution_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    analogy_box = doc.add_paragraph()
    analogy_box.add_run('Think of it like this:\n').bold = True
    analogy_box.add_run(
        'It\'s like planning a road trip to visit 20 cities. You wouldn\'t randomly jump '
        'from north to south to east! You\'d plan a route that visits nearby cities one after '
        'another, minimizing backtracking. That\'s exactly what this system does for sales routes.'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 2: WHAT THE SYSTEM DOES
    # ==========================
    doc.add_heading('2. What This System Does', 1)

    doc.add_heading('In Simple Terms:', 2)

    simple_steps = [
        ('Takes your customer list',
         'Reads all the customers each sales agent needs to visit'),

        ('Looks up their locations',
         'Finds where each customer is located (using GPS coordinates)'),

        ('Finds the shortest path',
         'Calculates the best route that connects all customers with minimal travel'),

        ('Assigns visit numbers',
         'Gives each customer a number: 1st stop, 2nd stop, 3rd stop, etc.'),

        ('Adds potential customers',
         'Suggests nearby prospects to visit while in the area'),

        ('Saves the plan',
         'Updates the database so agents can see their optimized route')
    ]

    for i, (step, explanation) in enumerate(simple_steps, start=1):
        doc.add_heading(f'Step {i}: {step}', 3)
        doc.add_paragraph(explanation, style='List Bullet 2')

    doc.add_paragraph()

    doc.add_heading('End Result:', 2)
    result_box = doc.add_paragraph()
    result_box.add_run('Every sales agent gets a numbered list showing:\n').bold = True
    result_box.add_run(
        '1st Stop: Customer A (closest to starting point)\n'
        '2nd Stop: Customer B (closest to Customer A)\n'
        '3rd Stop: Customer C (closest to Customer B)\n'
        '... and so on.\n\n'
        'This eliminates guesswork and wasted travel!'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 3: HOW IT WORKS
    # ==========================
    doc.add_heading('3. How It Works (The Journey)', 1)

    doc.add_paragraph(
        'Let\'s follow what happens when you run this system:'
    )

    doc.add_heading('Step 1: Organize by Groups', 2)
    doc.add_paragraph(
        'The system looks at your data and organizes it like a filing cabinet:'
    )
    org_example = doc.add_paragraph()
    org_example.add_run('📁 Distributor 1\n')
    org_example.add_run('  👤 Agent A\n')
    org_example.add_run('    📅 Monday, January 15\n')
    org_example.add_run('    📅 Tuesday, January 16\n')
    org_example.add_run('    📅 Wednesday, January 17\n')
    org_example.add_run('  👤 Agent B\n')
    org_example.add_run('    📅 Monday, January 15\n')
    org_example.add_run('    📅 Tuesday, January 16\n')

    doc.add_paragraph('This keeps everything neat and organized.')

    doc.add_paragraph()

    doc.add_heading('Step 2: Process Each Agent, One Day at a Time', 2)
    doc.add_paragraph('For each agent on each day, the system:')

    process_steps = [
        'Collects all customers assigned to that day',
        'Checks how many customers there are (is it a busy day or quiet day?)',
        'Gets the location (address) of each customer',
        'Identifies which customers have clear locations vs. missing information',
        'Decides if more customers should be added to fill the day'
    ]
    for step in process_steps:
        doc.add_paragraph(f'✓ {step}', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Step 3: Add Potential Customers (If Needed)', 2)

    target_box = doc.add_paragraph()
    target_box.add_run('Target: 60 visits per day\n\n').bold = True
    target_box.add_run(
        'Why 60? Research shows this is an efficient workload - enough to be productive '
        'without being overwhelming.\n\n'
        'If an agent only has 25 customers scheduled, the system will:\n'
        '• Look for potential new customers in the same neighborhood\n'
        '• Add up to 35 prospects to reach approximately 60 total visits\n'
        '• Only add prospects that make geographic sense (nearby locations)'
    )

    doc.add_paragraph()

    doc.add_heading('Step 4: Find the Shortest Route', 2)

    doc.add_paragraph(
        'This is where the magic happens! The system uses a smart method called '
        '"Nearest Neighbor" to find an efficient path:'
    )

    nn_steps = [
        'Start at the first customer (or a specified starting point)',
        'Look at all unvisited customers',
        'Choose the closest one as the next stop',
        'Repeat: From current location, go to the nearest unvisited customer',
        'Continue until all customers have been visited'
    ]

    for i, step in enumerate(nn_steps, start=1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph()

    example_box = doc.add_paragraph()
    example_box.add_run('Example:\n').bold = True
    example_box.add_run(
        'If you\'re at Customer A, and you have three unvisited customers:\n'
        '• Customer B is 2 km away\n'
        '• Customer C is 5 km away\n'
        '• Customer D is 8 km away\n\n'
        'The system picks Customer B (closest), then from B it finds the nearest customer, and so on.'
    )

    doc.add_paragraph()

    doc.add_heading('Step 5: Assign Visit Numbers', 2)

    doc.add_paragraph(
        'Once the best route is found, the system assigns numbers to each stop:'
    )

    numbering_para = doc.add_paragraph()
    numbering_para.add_run('Stop #1: First customer to visit\n')
    numbering_para.add_run('Stop #2: Second customer to visit\n')
    numbering_para.add_run('Stop #3: Third customer to visit\n')
    numbering_para.add_run('...\n')
    numbering_para.add_run('Stop #45: Last customer to visit\n\n')
    numbering_para.add_run('Special Case:\n').bold = True
    numbering_para.add_run(
        'If a customer has no address/location → Stop #100 (needs manual handling)'
    )

    doc.add_paragraph()

    doc.add_heading('Step 6: Save Everything', 2)

    doc.add_paragraph(
        'The system saves all these visit numbers to the database. Now when an agent '
        'logs in, they see their optimized route with numbered stops in order.'
    )

    doc.add_paragraph()

    doc.add_heading('Step 7: Repeat for Every Agent, Every Day', 2)

    doc.add_paragraph(
        'The system goes through every distributor, every agent, every date, and optimizes '
        'each one. This might take a few minutes depending on how much data you have.'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 4: DIFFERENT SITUATIONS
    # ==========================
    doc.add_heading('4. Different Situations (Scenarios)', 1)

    doc.add_paragraph(
        'The system handles all kinds of situations automatically. Here\'s what happens in different cases:'
    )

    doc.add_heading('Situation A: Busy Day (25+ customers)', 2)

    busy_table = doc.add_table(rows=4, cols=2)
    busy_table.style = 'Light Grid Accent 1'

    row = busy_table.rows[0].cells
    row[0].text = 'What happens:'
    row[1].text = 'Full optimization runs, huge time savings possible'

    row = busy_table.rows[1].cells
    row[0].text = 'Prospects added?'
    row[1].text = 'Yes, if under 60 total'

    row = busy_table.rows[2].cells
    row[0].text = 'Result:'
    row[1].text = 'Highly efficient route, minimal driving'

    row = busy_table.rows[3].cells
    row[0].text = 'Example:'
    row[1].text = '40 customers → add 20 prospects → 60 total stops, optimized'

    doc.add_paragraph()

    doc.add_heading('Situation B: Medium Day (10-24 customers)', 2)

    medium_table = doc.add_table(rows=4, cols=2)
    medium_table.style = 'Light Grid Accent 1'

    row = medium_table.rows[0].cells
    row[0].text = 'What happens:'
    row[1].text = 'Standard optimization, good improvements'

    row = medium_table.rows[1].cells
    row[0].text = 'Prospects added?'
    row[1].text = 'Yes, to reach 60'

    row = medium_table.rows[2].cells
    row[0].text = 'Result:'
    row[1].text = 'Balanced route with customers and prospects'

    row = medium_table.rows[3].cells
    row[0].text = 'Example:'
    row[1].text = '15 customers → add 45 prospects → 60 total stops'

    doc.add_paragraph()

    doc.add_heading('Situation C: Light Day (5-9 customers)', 2)

    light_table = doc.add_table(rows=4, cols=2)
    light_table.style = 'Light Grid Accent 1'

    row = light_table.rows[0].cells
    row[0].text = 'What happens:'
    row[1].text = 'Optimization still runs, day gets filled with prospects'

    row = light_table.rows[1].cells
    row[0].text = 'Prospects added?'
    row[1].text = 'Yes! Many prospects added'

    row = light_table.rows[2].cells
    row[0].text = 'Result:'
    row[1].text = 'Turns a slow day into a productive day'

    row = light_table.rows[3].cells
    row[0].text = 'Example:'
    row[1].text = '8 customers → add 52 prospects → 60 total stops'

    doc.add_paragraph()

    doc.add_heading('Situation D: Very Light Day (1-4 customers)', 2)

    vlight_table = doc.add_table(rows=4, cols=2)
    vlight_table.style = 'Light Grid Accent 1'

    row = vlight_table.rows[0].cells
    row[0].text = 'What happens:'
    row[1].text = 'Massively fills the day with prospects'

    row = vlight_table.rows[1].cells
    row[0].text = 'Prospects added?'
    row[1].text = 'Yes! Almost all stops will be prospects'

    row = vlight_table.rows[2].cells
    row[0].text = 'Result:'
    row[1].text = 'Makes the most of otherwise empty time'

    row = vlight_table.rows[3].cells
    row[0].text = 'Example:'
    row[1].text = '3 customers → add 57 prospects → 60 total stops'

    doc.add_page_break()

    # ==========================
    # SECTION 5: SPECIAL SITUATIONS
    # ==========================
    doc.add_heading('5. Special Situations Explained', 1)

    doc.add_heading('What if a customer has no address/location?', 2)

    no_address = doc.add_paragraph()
    no_address.add_run('Answer: ').bold = True
    no_address.add_run(
        'These customers get a special number: Stop #100\n\n'
        'This means: "This customer needs attention - please update their location or handle manually."\n\n'
        'They\'re not forgotten! They just can\'t be included in the automatic route because '
        'the system doesn\'t know where they are.'
    )

    doc.add_paragraph()

    doc.add_heading('What if there are no prospects available to add?', 2)

    no_prospects = doc.add_paragraph()
    no_prospects.add_run('Answer: ').bold = True
    no_prospects.add_run(
        'No problem! The system continues anyway.\n\n'
        'If an agent has 20 customers but there are no prospects in that area, the system will:\n'
        '• Optimize the route for those 20 customers\n'
        '• Assign stop numbers 1-20\n'
        '• Log a note saying "no prospects available"\n'
        '• Complete successfully\n\n'
        'The agent still gets an optimized route, just without the extra prospects.'
    )

    doc.add_paragraph()

    example_scenario = doc.add_paragraph()
    example_scenario.add_run('Real Example:\n').bold = True
    example_scenario.add_run(
        'Monday: Agent has 15 customers in Area A → System adds 45 prospects → 60 total\n'
        'Tuesday: Agent has 18 customers in Area A → No prospects left (all used Monday!) → 18 total\n'
        'Wednesday: Agent has 10 customers in Area B → System adds 50 prospects → 60 total\n\n'
        'Notice: Tuesday still worked fine even though prospects ran out!'
    )

    doc.add_paragraph()

    doc.add_heading('What if multiple agents need customers in the same area?', 2)

    same_area = doc.add_paragraph()
    same_area.add_run('Answer: ').bold = True
    same_area.add_run(
        'The system processes agents one at a time.\n\n'
        '• Agent 1 gets optimized first, prospects added\n'
        '• Agent 2 gets optimized next, gets remaining prospects\n'
        '• Agent 3 might find no prospects left\n\n'
        'Each prospect is only assigned once - no duplicates!\n'
        'The system checks to make sure a prospect isn\'t already assigned before adding it.'
    )

    doc.add_paragraph()

    doc.add_heading('Does each day start from Stop #1?', 2)

    numbering = doc.add_paragraph()
    numbering.add_run('Answer: YES! ').bold = True
    numbering.add_run(
        'Each day is independent.\n\n'
        'Monday: Stop 1, 2, 3... 45\n'
        'Tuesday: Stop 1, 2, 3... 38 (starts fresh)\n'
        'Wednesday: Stop 1, 2, 3... 52 (starts fresh)\n\n'
        'The stop number tells you the order to visit customers ON THAT SPECIFIC DAY, '
        'not across multiple days.'
    )

    doc.add_paragraph()

    doc.add_heading('Can I run this multiple times?', 2)

    multiple_runs = doc.add_paragraph()
    multiple_runs.add_run('Answer: Yes! ').bold = True
    multiple_runs.add_run(
        'Every time you run it, the system completely recalculates everything.\n\n'
        'It ignores old stop numbers and creates fresh ones. This is useful if:\n'
        '• Customer data changed\n'
        '• You added new customers\n'
        '• You want to re-optimize after updates'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 6: BENEFITS
    # ==========================
    doc.add_heading('6. Benefits of Using This System', 1)

    doc.add_heading('For Sales Agents:', 2)
    agent_benefits = [
        'Less driving = Less stress',
        'More time with customers, less time on the road',
        'Clear plan: No wondering "who should I visit next?"',
        'Visit more customers per day',
        'Save fuel and vehicle wear',
        'Get home earlier',
        'Discover new prospects along your route'
    ]
    for benefit in agent_benefits:
        doc.add_paragraph(f'✓ {benefit}', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('For Managers:', 2)
    manager_benefits = [
        'Better territory coverage',
        'More efficient use of resources',
        'Track which routes are busy vs. light',
        'Identify areas needing more/fewer agents',
        'Automatic prospect distribution',
        'Data-driven route planning',
        'Improved team productivity'
    ]
    for benefit in manager_benefits:
        doc.add_paragraph(f'✓ {benefit}', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('For the Company:', 2)
    company_benefits = [
        'Lower fuel costs',
        'Higher sales productivity',
        'Better customer service (agents not rushed)',
        'Reduced vehicle maintenance',
        'Environmental benefits (less driving)',
        'Competitive advantage (optimized operations)',
        'Scalable system (works for 10 agents or 1000 agents)'
    ]
    for benefit in company_benefits:
        doc.add_paragraph(f'✓ {benefit}', style='List Bullet')

    doc.add_page_break()

    # ==========================
    # SECTION 7: SIMPLE EXAMPLE
    # ==========================
    doc.add_heading('7. A Simple Example: Agent Maria\'s Day', 1)

    doc.add_heading('Before Optimization:', 2)

    before_para = doc.add_paragraph()
    before_para.add_run('Maria has 12 customers to visit on Friday.\n').bold = True
    before_para.add_run(
        'Without optimization, she might visit them in this order:\n\n'
        'Customer J → Customer A → Customer K → Customer C → Customer F...\n\n'
        'This could mean:\n'
        '• Driving north, then south, then north again (zigzagging!)\n'
        '• Total distance: 85 kilometers\n'
        '• Time on road: 4 hours\n'
        '• Fuel cost: $15\n'
        '• Stress level: High! 😰'
    )

    doc.add_paragraph()

    doc.add_heading('System Runs:', 2)

    process_para = doc.add_paragraph()
    process_para.add_run('The optimization system does its work:\n\n').bold = True
    process_para.add_run(
        '1️⃣ Looks at Maria\'s 12 customers\n'
        '2️⃣ Gets their locations from the database\n'
        '3️⃣ Finds 48 prospects in the same neighborhoods\n'
        '4️⃣ Calculates the shortest route connecting all 60 stops\n'
        '5️⃣ Assigns stop numbers 1-60\n'
        '6️⃣ Saves the optimized plan'
    )

    doc.add_paragraph()

    doc.add_heading('After Optimization:', 2)

    after_para = doc.add_paragraph()
    after_para.add_run('Maria now has an optimized route with 60 stops.\n').bold = True
    after_para.add_run(
        'The route flows naturally:\n\n'
        'Stop 1: Customer A (closest to starting point)\n'
        'Stop 2: Customer B (0.5 km from A)\n'
        'Stop 3: Prospect P123 (0.3 km from B)\n'
        'Stop 4: Customer C (0.4 km from P123)\n'
        '... and so on ...\n'
        'Stop 60: Customer L (on the way back)\n\n'
        'Results:\n'
        '• Total distance: 52 kilometers (38% reduction!)\n'
        '• Time on road: 2.5 hours (saved 1.5 hours)\n'
        '• Fuel cost: $9 (saved $6)\n'
        '• Visited: 60 customers/prospects instead of just 12!\n'
        '• Stress level: Low! Happy and productive! 😊'
    )

    doc.add_paragraph()

    impact_box = doc.add_paragraph()
    impact_box.add_run('The Impact:\n').bold = True
    impact_box.add_run(
        'Maria finishes her route faster, visits 5x more customers, uses less fuel, '
        'and has time to spare. She discovers new prospects she would have missed. '
        'This happens for EVERY agent, EVERY day!'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 8: HOW TO USE
    # ==========================
    doc.add_heading('8. How to Use This System', 1)

    doc.add_heading('Step 1: Prepare Your Data', 2)

    prep_para = doc.add_paragraph()
    prep_para.add_run('What you need:\n').bold = True
    prep_items = [
        'Customer assignments in the system (who visits whom)',
        'Customer locations (addresses or GPS coordinates)',
        'Prospect database with locations',
        'Route dates planned out'
    ]
    for item in prep_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    tip_box = doc.add_paragraph()
    tip_box.add_run('💡 Tip: ').bold = True
    tip_box.add_run(
        'The better your location data, the better the optimization! Make sure customer '
        'addresses are accurate and up to date.'
    )

    doc.add_paragraph()

    doc.add_heading('Step 2: Run the System', 2)

    run_para = doc.add_paragraph()
    run_para.add_run('Simple method:\n').bold = True
    run_para.add_run(
        '1. Open the system/program\n'
        '2. Click "Run Optimization" or execute the script\n'
        '3. Wait while it processes (usually 2-10 minutes)\n'
        '4. Check the log to see what happened\n'
        '5. Done!'
    )

    doc.add_paragraph()

    doc.add_heading('Step 3: Review Results', 2)

    review_para = doc.add_paragraph()
    review_para.add_run('Check the results:\n').bold = True
    review_items = [
        'Look at the log file to see how many routes were optimized',
        'Check for any warnings (like customers without addresses)',
        'Verify stop numbers were assigned',
        'Review any error messages',
        'Confirm all agents have their routes'
    ]
    for item in review_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Step 4: Agents Use the Routes', 2)

    use_para = doc.add_paragraph()
    use_para.add_run('Sales agents now see:\n').bold = True
    use_para.add_run(
        'A numbered list in their app/system showing:\n'
        '• Stop 1: Customer Name, Address\n'
        '• Stop 2: Customer Name, Address\n'
        '• Stop 3: Customer Name, Address\n'
        '... and so on\n\n'
        'They simply follow the numbers from 1 to 60 (or however many stops).\n'
        'No planning needed - just follow the sequence!'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 9: COMMON QUESTIONS
    # ==========================
    doc.add_heading('9. Common Questions (Simple Answers)', 1)

    questions = [
        ('How long does it take to run?',
         'Usually 2-10 minutes depending on how many agents and dates you have. '
         'For example, 50 agents × 20 days = 1000 routes might take 5 minutes.'),

        ('Does it change my customer data?',
         'No! It only reads customer information. The only thing it changes is the '
         'stop numbers in the route plan. Your original customer data stays exactly the same.'),

        ('What if something goes wrong?',
         'The system has safety features. If there\'s an error, it stops and doesn\'t make '
         'partial changes. Everything either completes successfully or nothing changes at all. '
         'Check the log file to see what happened.'),

        ('Can I change the route after optimization?',
         'Yes! The system provides a recommended route, but managers or agents can still '
         'adjust it manually if needed. The optimization is a helpful suggestion, not a strict rule.'),

        ('How often should I run it?',
         'Typically once per planning period (weekly or monthly). Run it whenever you:\n'
         '• Create new route assignments\n'
         '• Add new customers\n'
         '• Update customer locations\n'
         '• Want to refresh the routes'),

        ('Will it work if I have only a few agents?',
         'Absolutely! It works for any number of agents - whether you have 5 agents or 500 agents.'),

        ('What if two agents are in the same area?',
         'The system optimizes each agent independently. If two agents work the same area, '
         'they\'ll each get their own optimized route. Prospects won\'t be duplicated between them.'),

        ('How does it know what\'s "close"?',
         'It uses GPS coordinates (latitude/longitude) and calculates real distances between '
         'locations. It\'s the same way Google Maps knows distances - using the Earth\'s geography.'),

        ('Can I see the route on a map?',
         'This system calculates the stop numbers. To see the route on a map, you\'d export '
         'the data and use mapping software like Google Maps or your company\'s mapping tool.'),

        ('What happens to customers with no location?',
         'They get Stop #100, which is a special code meaning "needs manual handling." '
         'You should update their location and run the optimization again, or handle them manually.')
    ]

    for i, (question, answer) in enumerate(questions, start=1):
        doc.add_heading(f'Q{i}: {question}', 3)
        para = doc.add_paragraph(answer)
        para.runs[0].font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ==========================
    # SECTION 10: TROUBLESHOOTING
    # ==========================
    doc.add_heading('10. Simple Troubleshooting', 1)

    doc.add_heading('Problem: System says "No customers with coordinates"', 3)
    doc.add_paragraph('Solution:', style='Heading 4')
    sol_para = doc.add_paragraph()
    sol_para.add_run('Check your customer data. Make sure customers have addresses or GPS locations entered. '
                     'The system needs to know WHERE customers are to create routes. Update the missing locations '
                     'and run again.')

    doc.add_paragraph()

    doc.add_heading('Problem: System runs but no stop numbers appear', 3)
    doc.add_paragraph('Solution:', style='Heading 4')
    sol_para = doc.add_paragraph()
    sol_para.add_run('Check the log file for errors. There might be a database connection issue or a problem '
                     'with the data. Look for error messages in the log and share them with your IT team.')

    doc.add_paragraph()

    doc.add_heading('Problem: Some customers always get Stop #100', 3)
    doc.add_paragraph('Solution:', style='Heading 4')
    sol_para = doc.add_paragraph()
    sol_para.add_run('Stop #100 means "no location available." These customers are missing their GPS coordinates '
                     'or address. Update their information in the customer database, then re-run the optimization.')

    doc.add_paragraph()

    doc.add_heading('Problem: Takes too long to run', 3)
    doc.add_paragraph('Solution:', style='Heading 4')
    sol_para = doc.add_paragraph()
    sol_para.add_run('For very large datasets, you can:\n'
                     '• Run it during off-hours (overnight)\n'
                     '• Process one distributor at a time\n'
                     '• Ask IT to enable "parallel processing" mode (makes it faster)\n'
                     '• Check your database server performance')

    doc.add_paragraph()

    doc.add_heading('Problem: Different agents getting different results each time', 3)
    doc.add_paragraph('Solution:', style='Heading 4')
    sol_para = doc.add_paragraph()
    sol_para.add_run('This is normal when prospects are added. Prospects are selected randomly from available options. '
                     'The customer stops will stay consistent, but prospect assignments might vary. '
                     'This is intentional to ensure fair distribution.')

    doc.add_page_break()

    # ==========================
    # SECTION 11: SUMMARY
    # ==========================
    doc.add_heading('11. Summary: The Big Picture', 1)

    doc.add_heading('What You Need to Remember:', 2)

    summary_items = [
        ('Purpose',
         'This system helps sales agents visit customers in the most efficient order, '
         'saving time and increasing productivity.'),

        ('How it works',
         'It looks at customer locations, finds the shortest path connecting them, '
         'adds potential customers along the way, and assigns visit numbers.'),

        ('Main benefit',
         'Agents drive less, visit more customers, and have clear daily plans. '
         'The company saves money and operates more efficiently.'),

        ('What it changes',
         'Only the stop numbers (visit order). Customer data, addresses, and assignments '
         'remain unchanged.'),

        ('When to use it',
         'Run it whenever you create or update route plans - typically weekly or monthly.'),

        ('If something goes wrong',
         'Check the log file for messages. The system is designed to be safe - if there\'s '
         'an error, it won\'t make partial changes.'),

        ('Who benefits',
         'Everyone! Agents have easier days, managers have better oversight, '
         'customers get better service, and the company saves money.')
    ]

    for topic, explanation in summary_items:
        para = doc.add_paragraph()
        para.add_run(f'{topic}: ').bold = True
        para.add_run(explanation)

    doc.add_paragraph()

    final_box = doc.add_paragraph()
    final_box.add_run('Remember:\n').bold = True
    final_box.add_run(
        'You don\'t need to understand all the technical details. Just know that this system '
        'is like having an expert route planner working for you 24/7, finding the best paths '
        'for all your agents automatically. It turns complicated planning into simple numbered lists!\n\n'
    )
    final_box.add_run('Questions? ').bold = True
    final_box.add_run(
        'Contact your IT team or system administrator for help.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ==========================
    # VISUAL SUMMARY
    # ==========================
    doc.add_heading('Quick Visual Summary:', 2)

    visual = doc.add_paragraph()
    visual.add_run('INPUT ').bold = True
    visual.add_run('→ List of customers per agent per day\n')
    visual.add_run('         ↓\n')
    visual.add_run('PROCESS ').bold = True
    visual.add_run('→ Find locations, add prospects, calculate shortest route\n')
    visual.add_run('         ↓\n')
    visual.add_run('OUTPUT ').bold = True
    visual.add_run('→ Numbered list (Stop 1, Stop 2, Stop 3...)\n')
    visual.add_run('         ↓\n')
    visual.add_run('RESULT ').bold = True
    visual.add_run('→ Efficient routes, happy agents, lower costs!')

    # ==========================
    # FOOTER
    # ==========================
    doc.add_page_break()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('End of Guide')
    footer_run.font.size = Pt(14)
    footer_run.font.bold = True
    footer_run.font.color.rgb = RGBColor(0, 128, 0)

    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run2 = footer_para2.add_run('Thank you for reading!')
    footer_run2.font.size = Pt(11)

    footer_para3 = doc.add_paragraph()
    footer_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run3 = footer_para3.add_run(
        f'Document created: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
    )
    footer_run3.font.size = Pt(9)
    footer_run3.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = r'C:\Simplr projects\Route-optimization\full_pipeline\Route_Optimization_Simple_Guide.docx'
    doc.save(output_path)

    print(f"\n{'='*70}")
    print("SIMPLE GUIDE CREATED SUCCESSFULLY!")
    print(f"{'='*70}")
    print(f"\nSaved to: {output_path}")
    print(f"\nDocument Contents (Easy to Understand!):")
    print("  1. The Problem We're Solving (with examples)")
    print("  2. What This System Does (6 simple steps)")
    print("  3. How It Works - The Journey (detailed walkthrough)")
    print("  4. Different Situations (4 scenarios with tables)")
    print("  5. Special Situations Explained (5 common cases)")
    print("  6. Benefits (for agents, managers, company)")
    print("  7. A Simple Example - Maria's Day (before/after)")
    print("  8. How to Use This System (4 steps)")
    print("  9. Common Questions (10 Q&A in plain language)")
    print(" 10. Simple Troubleshooting (5 common problems)")
    print(" 11. Summary - The Big Picture (key takeaways)")
    print(f"\n{'='*70}")
    print("Written for NON-TECHNICAL users - no jargon!")
    print(f"{'='*70}")

if __name__ == "__main__":
    create_simple_overview()
