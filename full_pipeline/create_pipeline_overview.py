#!/usr/bin/env python3
"""
Create a high-level overview document for the Hierarchical Monthly Route Pipeline
Focusing on overall flow and scenario handling
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_overview():
    """Create overview document"""

    doc = Document()

    # Title
    title = doc.add_heading('Hierarchical Monthly Route Pipeline', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Overall Pipeline Flow & Scenario Handling')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(70, 130, 180)

    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================
    # SECTION 1: PIPELINE OVERVIEW
    # ==========================
    doc.add_heading('1. What Does This Pipeline Do?', 1)

    doc.add_paragraph(
        'This pipeline optimizes the order in which sales agents visit customers on their routes. '
        'It takes existing route plans and rearranges them to minimize travel distance, while also '
        'adding prospective customers to increase route efficiency.'
    )

    doc.add_heading('Think of it like this:', 2)
    analogy = [
        'Input: A list of customers each agent needs to visit on specific dates (but in random order)',
        'Process: Find the shortest path connecting all customers using geographic coordinates',
        'Output: An optimized sequence (StopNo 1, 2, 3...) that minimizes driving distance',
        'Bonus: Add nearby prospects to fill up the route (targeting 60 total stops per day)'
    ]
    for item in analogy:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ==========================
    # SECTION 2: HOW DATA FLOWS
    # ==========================
    doc.add_heading('2. How Data Flows Through The System', 1)

    doc.add_heading('The Big Picture:', 2)

    # Create visual flow
    flow_items = [
        ('START', 'MonthlyRoutePlan_temp table contains route assignments',
         'Fields: DistributorID, AgentID, RouteDate, CustNo, Name, etc.'),

        ('STEP 1: Build Hierarchy', 'System scans the table and creates a structure',
         'Distributor A → Agent 1, Agent 2...\n'
         'Distributor B → Agent 3, Agent 4...\n'
         'Each agent has multiple dates (chronologically sorted)'),

        ('STEP 2: Process by Hierarchy', 'Loop through: Distributor → Agent → Date',
         'This ensures agents are processed completely before moving to the next'),

        ('STEP 3: For Each Agent\'s Route Date', 'Perform optimization',
         'a) Get all customers assigned to this agent on this date\n'
         'b) Look up their coordinates (latitude/longitude) from customer table\n'
         'c) Separate: customers with coordinates vs. without\n'
         'd) Check how many customers we have (scenario classification)'),

        ('STEP 4: Enrich with Prospects', 'Add prospects if needed',
         'If total customers < 60:\n'
         '  - Find prospects in the same barangay (geographic area)\n'
         '  - Exclude prospects already visited or assigned\n'
         '  - Add enough to reach ~60 total stops\n'
         'If total customers >= 60: Skip prospect addition'),

        ('STEP 5: TSP Optimization', 'Find the shortest route',
         'Uses "Nearest Neighbor" algorithm:\n'
         '  - Start at first customer (or specified location)\n'
         '  - Always go to the nearest unvisited customer next\n'
         '  - Continue until all customers visited\n'
         '  - Assigns temporary StopNo: 1, 2, 3... N'),

        ('STEP 6: Assign Final StopNo', 'Create sequential stop numbers PER DATE',
         'For each date:\n'
         '  - Customers with coordinates: StopNo = 1, 2, 3... N (optimized order)\n'
         '  - Customers without coordinates: StopNo = 100\n'
         '  - Each date starts fresh at StopNo = 1'),

        ('STEP 7: Update Database', 'Save the optimized sequence',
         'UPDATE: Existing customers get new StopNo\n'
         'INSERT: New prospects added to MonthlyRoutePlan_temp\n'
         'All operations in a transaction (commit or rollback together)'),

        ('STEP 8: Next Iteration', 'Move to next date/agent/distributor',
         'Repeat steps 3-7 for every combination\n'
         'Progress logged continuously'),

        ('FINAL: Cleanup', 'Update customer type classification',
         'Mark records as "customer" or "prospect" using database JOINs'),

        ('END', 'Optimized routes in MonthlyRoutePlan_temp',
         'Each record has an optimized StopNo based on geographic proximity')
    ]

    for i, (stage, action, details) in enumerate(flow_items, start=1):
        doc.add_heading(f'{stage}', 3)
        doc.add_paragraph(action, style='Heading 4')

        # Split details by newline for better formatting
        for line in details.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip(), style='List Bullet 2')

        if i < len(flow_items):
            doc.add_paragraph('↓', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================
    # SECTION 3: PROCESSING ORDER
    # ==========================
    doc.add_heading('3. Why Process in This Order?', 1)

    doc.add_heading('Hierarchical Processing: Distributor → Agent → Date', 2)

    doc.add_paragraph(
        'The pipeline processes data in a strict hierarchy to ensure consistency and efficiency.'
    )

    order_table = doc.add_table(rows=4, cols=3)
    order_table.style = 'Medium Grid 3 Accent 1'

    header_cells = order_table.rows[0].cells
    header_cells[0].text = 'Level'
    header_cells[1].text = 'What It Means'
    header_cells[2].text = 'Why This Order?'

    order_data = [
        ('1. Distributor',
         'Process all agents under Distributor A, then move to Distributor B',
         'Keeps related data together, easier to track progress and handle errors'),

        ('2. Agent (within Distributor)',
         'Process all dates for Agent 1, then Agent 2, etc.',
         'Ensures sequential StopNo assignment across all dates for each agent'),

        ('3. Date (within Agent)',
         'Process dates chronologically (earliest to latest)',
         'Natural time order, easier to verify and debug')
    ]

    for i, (level, meaning, reason) in enumerate(order_data, start=1):
        cells = order_table.rows[i].cells
        cells[0].text = level
        cells[1].text = meaning
        cells[2].text = reason

    doc.add_paragraph()

    doc.add_heading('Per-Date StopNo Assignment (Key Concept)', 2)
    doc.add_paragraph(
        'IMPORTANT: Each date gets its own sequential StopNo starting from 1.'
    )

    example_box = doc.add_paragraph()
    example_box.add_run('Example:\n').bold = True
    example_box.add_run(
        'Agent 001 on 2025-01-15: StopNo 1, 2, 3... 45\n'
        'Agent 001 on 2025-01-16: StopNo 1, 2, 3... 38 (starts fresh)\n'
        'Agent 001 on 2025-01-17: StopNo 1, 2, 3... 52 (starts fresh)\n\n'
        'Each date is independent. StopNo represents the order of visits on THAT specific date.'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 4: SCENARIO HANDLING
    # ==========================
    doc.add_heading('4. Scenario Handling & Edge Cases', 1)

    doc.add_heading('4.1 Customer Count Scenarios', 2)

    scenario_table = doc.add_table(rows=5, cols=4)
    scenario_table.style = 'Light Grid Accent 1'

    header_cells = scenario_table.rows[0].cells
    header_cells[0].text = 'Scenario'
    header_cells[1].text = 'Customer Count'
    header_cells[2].text = 'What Happens'
    header_cells[3].text = 'Prospect Addition?'

    scenario_data = [
        ('Very Small Route', '1-4 customers',
         'TSP optimization still runs, minimal benefit',
         'Yes, add prospects to reach ~60'),

        ('Low Volume', '5-9 customers',
         'TSP provides modest improvements',
         'Yes, add prospects to reach ~60'),

        ('Medium Volume', '10-24 customers',
         'TSP provides good optimization',
         'Yes, add prospects to reach ~60'),

        ('High Volume', '25-59 customers',
         'TSP provides significant benefits',
         'Yes, fill remaining slots to ~60')
    ]

    for i, (scenario, count, happens, prospects) in enumerate(scenario_data, start=1):
        cells = scenario_table.rows[i].cells
        cells[0].text = scenario
        cells[1].text = count
        cells[2].text = happens
        cells[3].text = prospects

    doc.add_paragraph()

    note_para = doc.add_paragraph()
    note_para.add_run('Note: ').bold = True
    note_para.add_run(
        'If a route already has 60+ customers, no prospects are added. '
        'The system only adds prospects to routes with < 60 total customers.'
    )

    doc.add_heading('4.2 Coordinate Scenarios', 2)

    coord_scenarios = [
        ('Customers WITH coordinates',
         'These customers are included in TSP optimization and receive sequential StopNo (1, 2, 3...)'),

        ('Customers WITHOUT coordinates',
         'Assigned StopNo = 100 (special code meaning "no route optimization possible")'),

        ('Mix of both',
         'Customers with coordinates are optimized first, then customers without coordinates '
         'get StopNo = 100 at the end')
    ]

    for scenario, handling in coord_scenarios:
        doc.add_paragraph(f'{scenario}:', style='Heading 4')
        doc.add_paragraph(handling, style='List Bullet 2')

    doc.add_page_break()

    doc.add_heading('4.3 Prospect Addition Scenarios', 2)

    doc.add_paragraph(
        'This is where things get interesting. The system adds prospects intelligently '
        'based on geographic proximity and availability.'
    )

    doc.add_heading('Scenario A: Plenty of Prospects Available', 3)
    doc.add_paragraph(
        'Current customers: 20\n'
        'Needed prospects: 40 to reach 60\n'
        'Available prospects in same barangay: 150'
    )
    doc.add_paragraph('Result:', style='Heading 4')
    doc.add_paragraph('✓ System randomly selects 40 prospects from the barangay', style='List Bullet 2')
    doc.add_paragraph('✓ Prospects are added to this agent\'s route for this date', style='List Bullet 2')
    doc.add_paragraph('✓ Total route now has 60 stops', style='List Bullet 2')

    doc.add_heading('Scenario B: Limited Prospects Available', 3)
    doc.add_paragraph(
        'Current customers: 35\n'
        'Needed prospects: 25 to reach 60\n'
        'Available prospects in same barangay: 10'
    )
    doc.add_paragraph('Result:', style='Heading 4')
    doc.add_paragraph('✓ System adds all 10 available prospects', style='List Bullet 2')
    doc.add_paragraph('✓ Total route has 45 stops (not 60, but best possible)', style='List Bullet 2')
    doc.add_paragraph('✓ No error - system continues with what\'s available', style='List Bullet 2')

    doc.add_heading('Scenario C: All Prospects Already Assigned (Your Question!)', 3)
    doc.add_paragraph(
        'Current customers: 15\n'
        'Needed prospects: 45 to reach 60\n'
        'Available prospects in same barangay: 0 (all already assigned to previous agents/dates)'
    )
    doc.add_paragraph('Result:', style='Heading 4')
    doc.add_paragraph('✓ System tries to find prospects but finds none', style='List Bullet 2')
    doc.add_paragraph('✓ Logs: "No prospects found"', style='List Bullet 2')
    doc.add_paragraph('✓ Continues processing with just the 15 customers', style='List Bullet 2')
    doc.add_paragraph('✓ TSP optimization still runs on the 15 customers', style='List Bullet 2')
    doc.add_paragraph('✓ Route gets optimized StopNo even without reaching 60', style='List Bullet 2')

    doc.add_paragraph()
    important_box = doc.add_paragraph()
    important_box.add_run('IMPORTANT: ').bold = True
    important_box.add_run(
        'The pipeline does NOT fail or skip routes when prospects are unavailable. '
        'It simply processes whatever customers are available and logs the situation. '
        'This ensures all routes get optimized, even if prospect density targets can\'t be met.'
    )

    doc.add_heading('Scenario D: No Barangay Information', 3)
    doc.add_paragraph(
        'Current customers: 20, but none have coordinates or address3 (barangay_code)'
    )
    doc.add_paragraph('Result:', style='Heading 4')
    doc.add_paragraph('✗ System cannot determine which barangay to search for prospects', style='List Bullet 2')
    doc.add_paragraph('✓ Logs: "No barangay codes found - skipping prospect addition"', style='List Bullet 2')
    doc.add_paragraph('✓ Processes existing 20 customers only', style='List Bullet 2')
    doc.add_paragraph('✓ Customers without coordinates get StopNo = 100', style='List Bullet 2')

    doc.add_paragraph()
    note_box = doc.add_paragraph()
    note_box.add_run('Smart Design: ').bold = True
    note_box.add_run(
        'The system avoids adding random prospects when no geographic context is available. '
        'This prevents assigning prospects from distant locations that don\'t make sense for the route.'
    )

    doc.add_page_break()

    doc.add_heading('4.4 Prospect Exclusion Rules', 2)

    doc.add_paragraph('The system ensures prospects are not double-assigned by checking:')

    exclusion_table = doc.add_table(rows=4, cols=2)
    exclusion_table.style = 'Medium Grid 3 Accent 1'

    header_cells = exclusion_table.rows[0].cells
    header_cells[0].text = 'Check'
    header_cells[1].text = 'Purpose'

    exclusion_data = [
        ('NOT IN MonthlyRoutePlan_temp',
         'Ensures prospect isn\'t already assigned to this or another agent on this date'),

        ('NOT IN custvisit table',
         'Ensures prospect hasn\'t already been visited (converted to customer)'),

        ('Has valid coordinates',
         'Only prospects with latitude/longitude are eligible (needed for TSP)')
    ]

    for i, (check, purpose) in enumerate(exclusion_data, start=1):
        cells = exclusion_table.rows[i].cells
        cells[0].text = check
        cells[1].text = purpose

    doc.add_paragraph()

    doc.add_paragraph(
        'This means: Even if 1000 prospects exist in a barangay, only the ones that are:\n'
        '  • Not already assigned to any route\n'
        '  • Not already visited/converted\n'
        '  • Have valid coordinates\n'
        'will be considered for addition.'
    )

    doc.add_page_break()

    doc.add_heading('4.5 Database Update Scenarios', 2)

    doc.add_heading('Existing Customers (Already in MonthlyRoutePlan_temp)', 3)
    doc.add_paragraph('Action: UPDATE StopNo', style='List Bullet 2')
    doc.add_paragraph('SQL: UPDATE MonthlyRoutePlan_temp SET StopNo = ? WHERE ...', style='List Bullet 2')
    doc.add_paragraph('Result: Existing records get new optimized StopNo', style='List Bullet 2')

    doc.add_heading('New Prospects (Added for enrichment)', 3)
    doc.add_paragraph('Action: INSERT new record', style='List Bullet 2')
    doc.add_paragraph(
        'SQL: INSERT INTO MonthlyRoutePlan_temp (DistributorID, AgentID, RouteDate, CustNo, StopNo, ...)',
        style='List Bullet 2')
    doc.add_paragraph('Result: New records created with optimized StopNo', style='List Bullet 2')

    doc.add_heading('Transaction Safety', 3)
    doc.add_paragraph('All updates and inserts happen in a transaction', style='List Bullet 2')
    doc.add_paragraph('If any operation fails, everything rolls back', style='List Bullet 2')
    doc.add_paragraph('Ensures data integrity (all-or-nothing)', style='List Bullet 2')

    doc.add_page_break()

    # ==========================
    # SECTION 5: SPECIAL CASES
    # ==========================
    doc.add_heading('5. Special Cases & Edge Conditions', 1)

    special_cases = [
        ('Agent with No Dates',
         'If an agent has no dates in MonthlyRoutePlan_temp',
         'Skipped - logged as warning, processing continues'),

        ('Date with No Customers',
         'If a date has no valid customer records',
         'Skipped - logged as "no_data", processing continues'),

        ('All Customers Already Optimized',
         'If StopNo already exists (from previous run)',
         'IGNORED - existing StopNo is completely replaced with new values'),

        ('Customer Exists in Both customer AND prospective',
         'If CustNo found in multiple source tables',
         'Priority: customer table wins (custype = "customer")'),

        ('Duplicate CustNo on Same Date',
         'If same customer assigned twice to agent on same date',
         'Both records updated with same StopNo (database constraint should prevent this)'),

        ('Invalid Coordinates (0.0, NULL)',
         'If latitude/longitude is missing or invalid',
         'Customer gets StopNo = 100, excluded from TSP optimization'),

        ('TSP Optimization Fails',
         'If TSP algorithm encounters an error',
         'Logs error, continues with unoptimized order, assigns sequential StopNo'),

        ('Database Connection Lost',
         'If connection drops during processing',
         'Current transaction rolled back, error logged, pipeline stops'),

        ('Parallel Processing Enabled',
         'If --parallel flag used',
         'Multiple agents processed simultaneously, BUT dates within agent are still sequential'),

        ('Single Customer on Route',
         'If only 1 customer assigned to a date',
         'TSP optimization skipped (only 1 stop), StopNo = 1, prospect addition attempted')
    ]

    for i, (case, condition, handling) in enumerate(special_cases, start=1):
        doc.add_heading(f'{i}. {case}', 3)
        doc.add_paragraph(f'Condition: {condition}', style='List Bullet 2')
        doc.add_paragraph(f'Handling: {handling}', style='List Bullet 2')
        doc.add_paragraph()

    doc.add_page_break()

    # ==========================
    # SECTION 6: WHAT HAPPENS TO DATA
    # ==========================
    doc.add_heading('6. Data Transformation Summary', 1)

    doc.add_heading('Before Pipeline Runs:', 2)
    before_table = doc.add_table(rows=4, cols=5)
    before_table.style = 'Light List Accent 1'

    header_cells = before_table.rows[0].cells
    header_cells[0].text = 'DistributorID'
    header_cells[1].text = 'AgentID'
    header_cells[2].text = 'RouteDate'
    header_cells[3].text = 'CustNo'
    header_cells[4].text = 'StopNo'

    before_data = [
        ('D001', 'A101', '2025-01-15', 'C1234', 'NULL or random'),
        ('D001', 'A101', '2025-01-15', 'C5678', 'NULL or random'),
        ('D001', 'A101', '2025-01-15', 'C9012', 'NULL or random')
    ]

    for i, row_data in enumerate(before_data, start=1):
        cells = before_table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value

    doc.add_paragraph('↓ PIPELINE RUNS ↓', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('After Pipeline Runs:', 2)
    after_table = doc.add_table(rows=6, cols=6)
    after_table.style = 'Light Grid Accent 1'

    header_cells = after_table.rows[0].cells
    header_cells[0].text = 'DistributorID'
    header_cells[1].text = 'AgentID'
    header_cells[2].text = 'RouteDate'
    header_cells[3].text = 'CustNo'
    header_cells[4].text = 'StopNo'
    header_cells[5].text = 'custype'

    after_data = [
        ('D001', 'A101', '2025-01-15', 'C5678', '1', 'customer'),
        ('D001', 'A101', '2025-01-15', 'C9012', '2', 'customer'),
        ('D001', 'A101', '2025-01-15', 'P1111', '3', 'prospect'),
        ('D001', 'A101', '2025-01-15', 'C1234', '4', 'customer'),
        ('D001', 'A101', '2025-01-15', 'P2222', '5', 'prospect')
    ]

    for i, row_data in enumerate(after_data, start=1):
        cells = after_table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value

    doc.add_paragraph()

    doc.add_heading('What Changed:', 2)
    changes = [
        'StopNo now reflects geographic proximity (shortest path)',
        'Prospects (P1111, P2222) added to fill route',
        'custype column populated (customer vs. prospect)',
        'Order optimized: C5678 → C9012 → P1111 → C1234 → P2222 (nearest neighbor sequence)'
    ]
    for change in changes:
        doc.add_paragraph(change, style='List Bullet')

    doc.add_page_break()

    # ==========================
    # SECTION 7: PRACTICAL EXAMPLE
    # ==========================
    doc.add_heading('7. Complete Example Walkthrough', 1)

    doc.add_heading('Scenario: Process Agent A101 for Distributor D001', 2)

    doc.add_paragraph('Initial State:', style='Heading 3')
    doc.add_paragraph('Agent A101 has 3 dates assigned:', style='List Bullet 2')
    doc.add_paragraph('• 2025-01-15: 8 customers', style='List Bullet 3')
    doc.add_paragraph('• 2025-01-16: 15 customers', style='List Bullet 3')
    doc.add_paragraph('• 2025-01-17: 3 customers', style='List Bullet 3')

    doc.add_paragraph()

    doc.add_paragraph('Processing Flow:', style='Heading 3')

    doc.add_paragraph('DATE 1: 2025-01-15 (8 customers)', style='Heading 4')
    doc.add_paragraph('1. Check scenario: "Low Volume" (5-9 customers)', style='List Number 2')
    doc.add_paragraph('2. Get customer coordinates from database', style='List Number 2')
    doc.add_paragraph('   - 6 customers have valid coordinates', style='List Bullet 3')
    doc.add_paragraph('   - 2 customers have no coordinates', style='List Bullet 3')
    doc.add_paragraph('3. Identify barangay: All customers in "Brgy. San Jose"', style='List Number 2')
    doc.add_paragraph('4. Search for prospects in "Brgy. San Jose"', style='List Number 2')
    doc.add_paragraph('   - Need: 52 prospects (to reach 60 total)', style='List Bullet 3')
    doc.add_paragraph('   - Found: 45 available prospects', style='List Bullet 3')
    doc.add_paragraph('   - Added: All 45 prospects', style='List Bullet 3')
    doc.add_paragraph('5. Run TSP on 51 locations (6 customers + 45 prospects)', style='List Number 2')
    doc.add_paragraph('6. Assign StopNo:', style='List Number 2')
    doc.add_paragraph('   - Optimized 51 locations: StopNo 1-51', style='List Bullet 3')
    doc.add_paragraph('   - 2 without coordinates: StopNo 100', style='List Bullet 3')
    doc.add_paragraph('7. Update database:', style='List Number 2')
    doc.add_paragraph('   - UPDATE 8 existing customer records', style='List Bullet 3')
    doc.add_paragraph('   - INSERT 45 new prospect records', style='List Bullet 3')
    doc.add_paragraph('✓ Date 1 Complete: 53 total stops', style='Heading 4')

    doc.add_paragraph()

    doc.add_paragraph('DATE 2: 2025-01-16 (15 customers)', style='Heading 4')
    doc.add_paragraph('1. Check scenario: "Medium Volume" (10-24 customers)', style='List Number 2')
    doc.add_paragraph('2. Get customer coordinates', style='List Number 2')
    doc.add_paragraph('   - All 15 customers have valid coordinates', style='List Bullet 3')
    doc.add_paragraph('3. Identify barangay: "Brgy. San Jose"', style='List Number 2')
    doc.add_paragraph('4. Search for prospects', style='List Number 2')
    doc.add_paragraph('   - Need: 45 prospects', style='List Bullet 3')
    doc.add_paragraph('   - Found: 0 prospects (all 45 from this barangay were used on Date 1!)', style='List Bullet 3')
    doc.add_paragraph('   - Added: 0 prospects', style='List Bullet 3')
    doc.add_paragraph('   - Log: "No prospects found" (Warning, not Error)', style='List Bullet 3')
    doc.add_paragraph('5. Run TSP on 15 locations (customers only)', style='List Number 2')
    doc.add_paragraph('6. Assign StopNo: 1-15 (optimized order)', style='List Number 2')
    doc.add_paragraph('7. Update database:', style='List Number 2')
    doc.add_paragraph('   - UPDATE 15 existing customer records', style='List Bullet 3')
    doc.add_paragraph('   - INSERT 0 prospect records', style='List Bullet 3')
    doc.add_paragraph('✓ Date 2 Complete: 15 total stops (no prospects available, but route still optimized)', style='Heading 4')

    doc.add_paragraph()

    doc.add_paragraph('DATE 3: 2025-01-17 (3 customers)', style='Heading 4')
    doc.add_paragraph('1. Check scenario: "Very Small" (1-4 customers)', style='List Number 2')
    doc.add_paragraph('2. Get customer coordinates', style='List Number 2')
    doc.add_paragraph('   - 3 customers have valid coordinates', style='List Bullet 3')
    doc.add_paragraph('3. Identify barangay: "Brgy. Santa Cruz" (different from previous dates!)', style='List Number 2')
    doc.add_paragraph('4. Search for prospects in "Brgy. Santa Cruz"', style='List Number 2')
    doc.add_paragraph('   - Need: 57 prospects', style='List Bullet 3')
    doc.add_paragraph('   - Found: 120 available prospects', style='List Bullet 3')
    doc.add_paragraph('   - Added: 57 prospects (randomly selected)', style='List Bullet 3')
    doc.add_paragraph('5. Run TSP on 60 locations (3 customers + 57 prospects)', style='List Number 2')
    doc.add_paragraph('6. Assign StopNo: 1-60 (optimized order)', style='List Number 2')
    doc.add_paragraph('7. Update database:', style='List Number 2')
    doc.add_paragraph('   - UPDATE 3 existing customer records', style='List Bullet 3')
    doc.add_paragraph('   - INSERT 57 new prospect records', style='List Bullet 3')
    doc.add_paragraph('✓ Date 3 Complete: 60 total stops', style='Heading 4')

    doc.add_paragraph()

    doc.add_paragraph('Agent A101 Complete!', style='Heading 3')
    summary_box = doc.add_paragraph()
    summary_box.add_run('Summary:\n').bold = True
    summary_box.add_run(
        '• Total dates processed: 3\n'
        '• Total customer records updated: 26 (8 + 15 + 3)\n'
        '• Total prospect records inserted: 102 (45 + 0 + 57)\n'
        '• Total optimized stops: 128\n\n'
        'Key Insight: When prospects ran out on Date 2, the pipeline continued successfully '
        'with customer-only optimization. The route still received optimized StopNo values, '
        'just without the density bonus from prospects.'
    )

    doc.add_page_break()

    # ==========================
    # SECTION 8: KEY TAKEAWAYS
    # ==========================
    doc.add_heading('8. Key Takeaways', 1)

    doc.add_heading('The Pipeline Is Resilient:', 2)
    resilient_points = [
        'Works with any number of customers (1 to 1000+)',
        'Handles missing coordinates gracefully',
        'Continues when prospects are unavailable',
        'Processes routes even without reaching 60-stop target',
        'Logs warnings but doesn\'t fail on data issues',
        'Maintains transaction integrity (rollback on errors)'
    ]
    for point in resilient_points:
        doc.add_paragraph(f'✓ {point}', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('The Pipeline Is Intelligent:', 2)
    intelligent_points = [
        'Only adds prospects from the same geographic area (barangay)',
        'Avoids double-assigning prospects to multiple agents',
        'Skips prospects that have already been visited',
        'Uses efficient TSP algorithm for route optimization',
        'Assigns meaningful StopNo = 100 for unoptimizable customers',
        'Maintains separate numbering per date (not cumulative)'
    ]
    for point in intelligent_points:
        doc.add_paragraph(f'✓ {point}', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('The Pipeline Is Transparent:', 2)
    transparent_points = [
        'Comprehensive logging of every step',
        'Progress tracking (X of Y combinations processed)',
        'Scenario classification for each route',
        'Detailed error messages when issues occur',
        'Final summary with success/error/skip counts',
        'Timestamped log files for audit trail'
    ]
    for point in transparent_points:
        doc.add_paragraph(f'✓ {point}', style='List Bullet')

    doc.add_page_break()

    # ==========================
    # SECTION 9: VISUAL SUMMARY
    # ==========================
    doc.add_heading('9. Visual Summary', 1)

    doc.add_heading('Simple Mental Model:', 2)

    visual_flow = [
        'INPUT',
        '  ↓',
        'Read MonthlyRoutePlan_temp',
        '  ↓',
        'Build Hierarchy (Distributor → Agent → Date)',
        '  ↓',
        'FOR EACH Distributor:',
        '    FOR EACH Agent:',
        '        FOR EACH Date (chronological):',
        '            ↓',
        '            Get Customers',
        '            ↓',
        '            Add Prospects (if needed & available)',
        '            ↓',
        '            Optimize Route (TSP)',
        '            ↓',
        '            Assign StopNo (1, 2, 3... or 100)',
        '            ↓',
        '            Save to Database',
        '  ↓',
        'Update Customer Types',
        '  ↓',
        'OUTPUT',
        '  ↓',
        'Optimized Routes in MonthlyRoutePlan_temp'
    ]

    for line in visual_flow:
        para = doc.add_paragraph(line)
        if line in ['INPUT', 'OUTPUT']:
            para.runs[0].bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif '→' in line or 'FOR EACH' in line:
            para.runs[0].font.color.rgb = RGBColor(0, 100, 200)

    doc.add_page_break()

    # ==========================
    # SECTION 10: ANSWERS TO COMMON QUESTIONS
    # ==========================
    doc.add_heading('10. Answers to Common Questions', 1)

    qa_pairs = [
        ('What if all prospects are already assigned to previous agents?',
         'The system logs "No prospects found" and continues processing with just the existing '
         'customers. The route still gets optimized and receives sequential StopNo values. '
         'This is a warning, not an error.'),

        ('Does each date start StopNo from 1?',
         'YES! Each RouteDate gets its own independent StopNo sequence starting from 1. '
         'This is intentional - StopNo represents the visit order on THAT specific date.'),

        ('What happens to customers without coordinates?',
         'They receive StopNo = 100 as a special marker. They are excluded from TSP optimization '
         'but remain in the route plan. This signals they need manual handling or coordinate updates.'),

        ('Can prospects be assigned to multiple agents on the same date?',
         'NO. The prospect exclusion query checks MonthlyRoutePlan_temp to prevent double assignment. '
         'Once a prospect is assigned to any agent on a date, it becomes unavailable for other agents '
         'on that date.'),

        ('How does the system choose which prospects to add?',
         'Randomly (ORDER BY NEWID()) from eligible prospects in the same barangay. This ensures '
         'fair distribution and geographic relevance.'),

        ('What if a customer appears in both customer and prospective tables?',
         'The customer table takes priority. The custype will be set to "customer" not "prospect".'),

        ('Does the pipeline modify the customer or prospective tables?',
         'NO. It only reads from these tables. All modifications are made to MonthlyRoutePlan_temp only.'),

        ('Can I run the pipeline multiple times?',
         'YES. Each run completely overwrites previous StopNo values. This allows re-optimization '
         'if data changes or if you want to try different settings.'),

        ('What if the database connection fails mid-process?',
         'The current transaction rolls back, ensuring no partial updates. The error is logged, '
         'and the pipeline stops. You can re-run after fixing the connection issue.'),

        ('How long does it take to process?',
         'Depends on data volume. Typical rate: 5-20 combinations/second. For example, 1000 '
         'combinations might take 1-3 minutes. Use --parallel flag for faster processing.')
    ]

    for i, (question, answer) in enumerate(qa_pairs, start=1):
        doc.add_heading(f'Q{i}: {question}', 3)
        doc.add_paragraph(answer, style='List Bullet 2')
        doc.add_paragraph()

    # ==========================
    # FOOTER
    # ==========================
    doc.add_page_break()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('End of Pipeline Overview')
    footer_run.font.size = Pt(12)
    footer_run.font.bold = True
    footer_run.font.color.rgb = RGBColor(70, 130, 180)

    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run2 = footer_para2.add_run(
        f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
    )
    footer_run2.font.size = Pt(9)
    footer_run2.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = r'C:\Simplr projects\Route-optimization\full_pipeline\Pipeline_Overview_And_Scenarios.docx'
    doc.save(output_path)

    print(f"\n{'='*70}")
    print("OVERVIEW DOCUMENT CREATED SUCCESSFULLY!")
    print(f"{'='*70}")
    print(f"\nSaved to: {output_path}")
    print(f"\nDocument Contents:")
    print("  1. What Does This Pipeline Do?")
    print("  2. How Data Flows Through The System (detailed flow)")
    print("  3. Why Process in This Order? (hierarchy explanation)")
    print("  4. Scenario Handling & Edge Cases")
    print("     - Customer count scenarios")
    print("     - Coordinate scenarios")
    print("     - Prospect addition scenarios (INCLUDING your question!)")
    print("     - Prospect exclusion rules")
    print("     - Database update scenarios")
    print("  5. Special Cases & Edge Conditions (10 cases)")
    print("  6. Data Transformation Summary (before/after)")
    print("  7. Complete Example Walkthrough")
    print("     - 3 dates processed with different scenarios")
    print("     - Shows what happens when prospects run out!")
    print("  8. Key Takeaways (Resilient, Intelligent, Transparent)")
    print("  9. Visual Summary (mental model)")
    print(" 10. Answers to Common Questions (10 Q&A pairs)")
    print(f"\n{'='*70}")

if __name__ == "__main__":
    create_overview()
